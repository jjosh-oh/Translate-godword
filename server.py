"""
Translation server with operator dashboard.
"""
import os
import io
import json
import queue
import threading
import anthropic
from flask import Flask, request, jsonify, Response, send_from_directory
from flask_sock import Sock
from dotenv import load_dotenv

load_dotenv()

# Google Cloud 인증: 폴더 내 google-key.json 자동 사용
_key_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "google-key.json")
if os.path.exists(_key_path) and "GOOGLE_APPLICATION_CREDENTIALS" not in os.environ:
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = _key_path

app = Flask(__name__)
sock = Sock(app)
client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

# Gemini 클라이언트(비교용) — Vertex AI 경유, 실패해도 서버는 정상 동작
GEMINI_PROJECT = "project-71ecd8e1-9699-417f-ae4"
gemini_client = None
try:
    from google import genai as _genai
    gemini_client = _genai.Client(vertexai=True, project=GEMINI_PROJECT, location="us-central1")
except Exception as _e:
    print("Gemini 초기화 건너뜀:", _e)

# 음성인식 언어 코드 매핑 (Google Speech 형식)
STT_LANG = {
    "ko-KR": "ko-KR", "en-US": "en-US", "ja-JP": "ja-JP",
    "zh-CN": "cmn-Hans-CN", "es-ES": "es-ES", "fr-FR": "fr-FR", "de-DE": "de-DE",
}
STT_TO_NAME = {
    "ko-KR": "Korean", "en-US": "English", "ja-JP": "Japanese",
    "zh-CN": "Chinese", "es-ES": "Spanish", "fr-FR": "French", "de-DE": "German",
}

# 여러 기기(송출창 + 셀폰들)에 동시 전송하기 위한 구독자 기반 브로드캐스트
_display_subscribers = set()
_operator_subscribers = set()
_sub_lock = threading.Lock()


class _Broadcaster:
    """put() 호출 시 모든 구독자 큐에 메시지를 복사해 넣는다."""
    def __init__(self, subscribers):
        self._subs = subscribers

    def put(self, item):
        with _sub_lock:
            for q in list(self._subs):
                q.put(item)

    def subscribe(self):
        q = queue.Queue()
        with _sub_lock:
            self._subs.add(q)
        return q

    def unsubscribe(self, q):
        with _sub_lock:
            self._subs.discard(q)


translation_queue = _Broadcaster(_display_subscribers)
operator_queue = _Broadcaster(_operator_subscribers)

# Shared display settings
settings = {
    "font_size": 5,
    "text_color": "#ffffff",
    "bg_color": "#000000",
    "delay": 0,
    "target_lang": "English",
    "source_lang": "Korean",
}

last_input = ""
last_output = ""
sermon_context = ""  # 업로드된 설교 자료

# 교회 용어집(인식 힌트 전용) — glossary.txt에서 한 줄에 하나씩 로드
glossary_terms = []


def load_glossary():
    global glossary_terms
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "glossary.txt")
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            lines = [ln.strip() for ln in f.read().splitlines()]
        glossary_terms = [ln for ln in lines if ln]
    except Exception:
        glossary_terms = []


load_glossary()

# 번역 대응표 (한글=English 형식)
translation_mapping = {}

def load_mapping():
    global translation_mapping
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "mapping.txt")
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    line = line.strip()
                    if "=" in line and line:
                        k, v = line.split("=", 1)
                        translation_mapping[k.strip()] = v.strip()
    except Exception:
        translation_mapping = {}

load_mapping()


def _log_translation(src, out):
    try:
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "로그.txt")
        import datetime
        ts = datetime.datetime.now().strftime("%H:%M:%S")
        with open(path, "a", encoding="utf-8") as f:
            f.write(f"[{ts}] 입력: {src}\n[{ts}] 번역: {out}\n\n")
    except Exception:
        pass


def translate_and_stream(text: str, target_lang: str, source_lang: str):
    global last_output
    base_instruction = (
        f"You are a professional church interpreter providing live subtitles. "
        f"Translate ONLY the exact given text from {source_lang} into {target_lang}. "
        f"Output ONLY the translation of what is given — never add, continue, complete, "
        f"or quote additional text. If the input is a short or partial sentence (e.g. part of "
        f"a Bible verse), translate only that fragment; do NOT finish the verse or add the rest. "
        f"No explanations or commentary."
    )

    # 번역 대응표가 있으면 프롬프트에 추가
    mapping_text = ""
    if translation_mapping:
        pairs = "\n".join(f"  {k} → {v}" for k, v in translation_mapping.items())
        mapping_text = (
            "\n\nName/term translation table (use these EXACT translations when the term appears):\n"
            + pairs
        )

    if sermon_context:
        # 설교 원고: 고유명사 확인 '참고용'으로만 — 내용을 이어 쓰지 않도록 명시
        reference = (
            "The sermon script below is a REFERENCE ONLY, used solely to spell proper nouns "
            "(biblical names, place names, Korean locations) consistently. "
            "Do NOT translate, output, continue, or copy any sentence from this script. "
            "Only translate the user's given text.\n\n"
            "--- SERMON REFERENCE (do not output) ---\n" + sermon_context[:8000] + "\n--- END ---"
        )
        system = [
            {"type": "text", "text": base_instruction + mapping_text},
            {"type": "text", "text": reference,
             "cache_control": {"type": "ephemeral"}},
        ]
    else:
        system = base_instruction + mapping_text

    translation_queue.put("__clear__")
    operator_queue.put(("output_clear", ""))

    # 출력 폭주 방지: 입력 길이에 비례한 상한(최소 256, 최대 1024 토큰)
    max_out = max(256, min(1024, len(text) * 4))

    result = []
    with client.messages.stream(
        model="claude-opus-4-8",
        max_tokens=max_out,
        system=system,
        messages=[{"role": "user", "content": text}],
    ) as stream:
        for chunk in stream.text_stream:
            result.append(chunk)
            translation_queue.put(chunk)
            operator_queue.put(("output_chunk", chunk))

    last_output = "".join(result)
    _log_translation(text, last_output)
    translation_queue.put("__done__")
    operator_queue.put(("done", ""))


# ===== 번역 작업 큐 (문장 단위 번역을 순서대로 처리) =====
translation_jobs = queue.Queue()

# 음성 세션 관리: 마이크를 끄면 세션이 바뀌어 대기 중 번역이 폐기됨
_current_session = 0
_session_lock = threading.Lock()


def _translation_worker():
    while True:
        text, target_lang, source_lang, session = translation_jobs.get()
        # 세션이 유효한 경우에만 번역 (None이면 수동 입력 → 항상 실행)
        if session is not None:
            with _session_lock:
                if session != _current_session:
                    continue  # 마이크가 꺼진 뒤의 오래된 작업 → 폐기
        try:
            translate_and_stream(text, target_lang, source_lang)
        except Exception as e:
            print("번역 오류:", e)


threading.Thread(target=_translation_worker, daemon=True).start()


def enqueue_translation(text, target_lang, source_lang, session=None):
    translation_jobs.put((text, target_lang, source_lang, session))


@app.route("/")
def index():
    return send_from_directory(".", "display.html")


@app.route("/operator")
def operator():
    return send_from_directory(".", "operator.html")


@app.route("/mobile")
def mobile():
    return send_from_directory(".", "mobile.html")


@app.route("/compare")
def compare():
    return send_from_directory(".", "compare.html")


def _translate_claude(text, target_lang, source_lang):
    system = (f"You are a professional church interpreter. Translate from {source_lang} into {target_lang}. "
              f"Output ONLY the translation.")
    m = client.messages.create(model="claude-opus-4-8", max_tokens=1024,
                               system=system, messages=[{"role": "user", "content": text}])
    return "".join(b.text for b in m.content if getattr(b, "type", "") == "text")


def _translate_gemini(text, target_lang, source_lang):
    if not gemini_client:
        return "(Gemini 사용 불가)"
    prompt = (f"You are a professional church interpreter. Translate from {source_lang} into {target_lang}. "
              f"Output ONLY the translation.\n\n{text}")
    r = gemini_client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
    return (r.text or "").strip()


@sock.route("/audio-compare")
def audio_compare(ws):
    """비교용: 음성인식 후 같은 문장을 Claude·Gemini로 번역해 나란히 전송."""
    from google.cloud import speech
    import time as _t

    cfg = json.loads(ws.receive())
    src_code = STT_LANG.get(cfg.get("source_lang_code"), "ko-KR")
    source_name = STT_TO_NAME.get(cfg.get("source_lang_code"), "Korean")
    target_lang = cfg.get("target_lang", "English")

    speech_client = speech.SpeechClient()
    recog_config = speech.RecognitionConfig(
        encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=16000, language_code=src_code,
        enable_automatic_punctuation=True, model="latest_long", use_enhanced=True,
    )
    streaming_config = speech.StreamingRecognitionConfig(config=recog_config, interim_results=True)

    audio_q = queue.Queue()
    stop_flag = {"stop": False}
    send_lock = threading.Lock()

    def safe_send(obj):
        try:
            with send_lock:
                ws.send(json.dumps(obj))
        except Exception:
            pass

    def receiver():
        try:
            while True:
                data = ws.receive()
                if data is None:
                    break
                if isinstance(data, (bytes, bytearray)):
                    audio_q.put(bytes(data))
        except Exception:
            pass
        finally:
            stop_flag["stop"] = True
            audio_q.put(None)

    threading.Thread(target=receiver, daemon=True).start()

    def request_gen():
        while not stop_flag["stop"]:
            chunk = audio_q.get()
            if chunk is None:
                return
            yield speech.StreamingRecognizeRequest(audio_content=chunk)

    def run_engine(engine, fn, text):
        t0 = _t.time()
        try:
            out = fn(text, target_lang, source_name)
        except Exception as e:
            out = "(오류: " + str(e)[:60] + ")"
        ms = int((_t.time() - t0) * 1000)
        safe_send({"engine": engine, "text": out, "ms": ms})

    import re as _rec

    def split_sentences(t):
        return [p for p in _rec.split(r'(?<=[\.\?\!。？！…])\s*', t) if p.strip()]

    def translate_both(text):
        safe_send({"engine": "input", "text": text})
        threading.Thread(target=run_engine, args=("claude", _translate_claude, text), daemon=True).start()
        threading.Thread(target=run_engine, args=("gemini", _translate_gemini, text), daemon=True).start()

    translated = []

    while not stop_flag["stop"]:
        try:
            responses = speech_client.streaming_recognize(streaming_config, request_gen())
            for response in responses:
                for result in response.results:
                    tr = result.alternatives[0].transcript
                    if not tr.strip():
                        continue
                    pieces = split_sentences(tr)
                    if result.is_final:
                        for p in pieces:
                            n = p.strip()
                            if n and n not in translated:
                                translate_both(n)
                        translated = []
                    else:
                        safe_send({"engine": "interim", "text": tr.strip()})
                        if len(pieces) >= 2:
                            for p in pieces[:-1]:
                                n = p.strip()
                                if n and n not in translated:
                                    translated.append(n)
                                    translate_both(n)
        except Exception as e:
            safe_send({"engine": "error", "text": str(e)[:120]})
            if stop_flag["stop"]:
                break


@app.route("/tunnel-url")
def tunnel_url():
    """공개 주소 반환. ngrok 고정 도메인이 설정돼 있으면 우선 사용, 없으면 cloudflare 임시주소."""
    import re
    base = os.path.dirname(os.path.abspath(__file__))

    # 1) ngrok 고정 도메인 우선
    domain_path = os.path.join(base, "ngrok-domain.txt")
    try:
        if os.path.exists(domain_path):
            with open(domain_path, "r", encoding="utf-8", errors="ignore") as f:
                domain = f.read().strip()
            if domain:
                return jsonify({"url": "https://" + domain})
    except Exception:
        pass

    # 2) cloudflare 임시 주소(tunnel.log)
    log_path = os.path.join(base, "tunnel.log")
    try:
        with open(log_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        matches = re.findall(r"https://[a-zA-Z0-9-]+\.trycloudflare\.com", text)
        if matches:
            return jsonify({"url": matches[-1]})
    except Exception:
        pass
    return jsonify({"url": ""})


@sock.route("/audio")
def audio_socket(ws):
    """브라우저에서 16kHz PCM 오디오를 받아 Google Speech로 스트리밍 인식 → 번역."""
    from google.cloud import speech

    # 첫 메시지: 설정(JSON)
    cfg_raw = ws.receive()
    cfg = json.loads(cfg_raw)
    src_code = STT_LANG.get(cfg.get("source_lang_code"), "ko-KR")
    source_name = STT_TO_NAME.get(cfg.get("source_lang_code"), "Korean")
    target_lang = cfg.get("target_lang", "English")

    # 인식 힌트(speech adaptation)
    speech_contexts = []
    # 1) 교회 용어집 — 모든 항목을 높은 가중치로(고유명사 사전)
    if glossary_terms:
        speech_contexts.append(speech.SpeechContext(phrases=glossary_terms[:4000], boost=20.0))
    # 2) 오늘 설교 원고에서 자주 나오는 단어 — 보조 힌트
    if sermon_context:
        import re as _re2
        from collections import Counter
        words = _re2.findall(r"[가-힣]{2,}", sermon_context)
        common = [w for w, _ in Counter(words).most_common(300)]
        if common:
            speech_contexts.append(speech.SpeechContext(phrases=common, boost=12.0))

    speech_client = speech.SpeechClient()
    recog_config = speech.RecognitionConfig(
        encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=16000,
        language_code=src_code,
        enable_automatic_punctuation=True,
        model="latest_long",        # 긴 발화에 적합한 최신 모델
        use_enhanced=True,          # 고품질(enhanced) 모델 사용
        speech_contexts=speech_contexts,
    )
    streaming_config = speech.StreamingRecognitionConfig(
        config=recog_config, interim_results=True
    )

    audio_q = queue.Queue()
    stop_flag = {"stop": False}

    # 브라우저 → 오디오 수신 스레드
    def receiver():
        try:
            while True:
                data = ws.receive()
                if data is None:
                    break
                if isinstance(data, (bytes, bytearray)):
                    audio_q.put(bytes(data))
        except Exception:
            pass
        finally:
            stop_flag["stop"] = True
            audio_q.put(None)

    recv_thread = threading.Thread(target=receiver, daemon=True)
    recv_thread.start()

    def request_gen():
        # Google 스트림은 약 5분 제한 → 호출부에서 재시작
        while not stop_flag["stop"]:
            chunk = audio_q.get()
            if chunk is None:
                return
            yield speech.StreamingRecognizeRequest(audio_content=chunk)

    # 이 연결의 세션 ID (마이크를 끄면 무효화되어 대기 중 번역이 폐기됨)
    global _current_session
    with _session_lock:
        _current_session += 1
        my_session = _current_session

    import re as _re

    def split_sentences(t):
        # 문장부호(. ? ! 。 ？ ！ …) 뒤에서 자른다. 부호를 포함해 반환.
        parts = _re.split(r'(?<=[\.\?\!。？！…])\s*', t)
        return [p for p in parts if p.strip()]

    translated = []  # 현재 발화에서 이미 번역에 보낸 문장(텍스트로 중복 제거)

    # 5분 제한 대응: 스트림이 끝나면 다시 연결 (음성 큐는 유지)
    while not stop_flag["stop"]:
        try:
            responses = speech_client.streaming_recognize(streaming_config, request_gen())
            for response in responses:
                for result in response.results:
                    transcript = result.alternatives[0].transcript
                    if not transcript.strip():
                        continue
                    pieces = split_sentences(transcript)

                    if result.is_final:
                        operator_queue.put(("input", transcript.strip()))
                        # 아직 안 보낸 문장(끝 미완성 조각 포함)을 모두 번역
                        for p in pieces:
                            n = p.strip()
                            if n and n not in translated:
                                enqueue_translation(n, target_lang, source_name, my_session)
                        translated = []  # 다음 발화 위해 초기화
                    else:
                        operator_queue.put(("interim", transcript.strip()))
                        # 마침표로 '완성된' 문장만 즉시 번역 (마지막 미완성 조각 제외)
                        if len(pieces) >= 2:
                            complete = pieces[:-1]
                            for p in complete:
                                n = p.strip()
                                if n and n not in translated:
                                    translated.append(n)
                                    enqueue_translation(n, target_lang, source_name, my_session)
        except Exception as e:
            operator_queue.put(("stt_error", str(e)[:120]))
            if stop_flag["stop"]:
                break

    # 마이크 종료: 이 세션을 무효화해 대기 중인 번역 작업을 폐기
    with _session_lock:
        if _current_session == my_session:
            _current_session += 1


@app.route("/translate", methods=["POST"])
def translate():
    global last_input
    data = request.json
    text = data.get("text", "").strip()
    target_lang = data.get("target_lang", settings["target_lang"])
    source_lang = data.get("source_lang", settings["source_lang"])

    if not text:
        return jsonify({"error": "텍스트를 입력하세요."}), 400

    last_input = text
    operator_queue.put(("input", text))

    thread = threading.Thread(target=translate_and_stream, args=(text, target_lang, source_lang))
    thread.daemon = True
    thread.start()

    return jsonify({"status": "started"})


@app.route("/upload", methods=["POST"])
def upload():
    global sermon_context
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "파일이 없습니다."}), 400

    filename = file.filename.lower()
    try:
        if filename.endswith(".txt"):
            sermon_context = file.read().decode("utf-8", errors="ignore")

        elif filename.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file.read()))
            sermon_context = "\n".join(page.extract_text() or "" for page in reader.pages)

        elif filename.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(file.read()))
            sermon_context = "\n".join(p.text for p in doc.paragraphs)

        else:
            return jsonify({"error": "지원 형식: .txt, .pdf, .docx"}), 400

        preview = sermon_context[:200].replace("\n", " ")
        return jsonify({"status": "ok", "chars": len(sermon_context), "preview": preview})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/upload-clear", methods=["POST"])
def upload_clear():
    global sermon_context
    sermon_context = ""
    return jsonify({"status": "cleared"})


def _extract_text(file):
    name = file.filename.lower()
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file.read()))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if name.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(file.read()))
        return "\n".join(p.text for p in doc.paragraphs)
    return file.read().decode("utf-8", errors="ignore")


@app.route("/upload-glossary", methods=["POST"])
def upload_glossary():
    """교회 용어집 업로드 — 기존 용어집에 병합(중복 제거)하여 저장."""
    global glossary_terms
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "파일이 없습니다."}), 400
    try:
        text = _extract_text(file)
        # 기존 용어 유지 + 새 용어 추가 (순서 보존, 중복 제거)
        seen = set(glossary_terms)
        merged = list(glossary_terms)
        new_count = 0
        for ln in text.splitlines():
            t = ln.strip()
            if t and t not in seen:
                seen.add(t)
                merged.append(t)
                new_count += 1
        glossary_terms = merged
        # glossary.txt에 저장(다음 실행에도 유지)
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "glossary.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(merged))
        return jsonify({"status": "ok", "count": len(merged), "added": new_count,
                        "preview": ", ".join(merged[:15])})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/glossary-info")
def glossary_info():
    return jsonify({"count": len(glossary_terms),
                    "preview": ", ".join(glossary_terms[:15])})


@app.route("/upload-mapping", methods=["POST"])
def upload_mapping():
    """번역 대응표 업로드 — 한글=English 형식, 기존 항목에 병합."""
    global translation_mapping
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "파일이 없습니다."}), 400
    try:
        text = _extract_text(file)
        new_count = 0
        for line in text.splitlines():
            line = line.strip()
            if "=" in line and line:
                k, v = line.split("=", 1)
                k, v = k.strip(), v.strip()
                if k and v:
                    if k not in translation_mapping:
                        new_count += 1
                    translation_mapping[k] = v
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "mapping.txt")
        with open(path, "w", encoding="utf-8") as f:
            for k, v in translation_mapping.items():
                f.write(f"{k}={v}\n")
        preview = ", ".join(f"{k}→{v}" for k, v in list(translation_mapping.items())[:5])
        return jsonify({"status": "ok", "count": len(translation_mapping),
                        "added": new_count, "preview": preview})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/mapping-info")
def mapping_info():
    preview = ", ".join(f"{k}→{v}" for k, v in list(translation_mapping.items())[:5])
    return jsonify({"count": len(translation_mapping), "preview": preview})


@app.route("/settings", methods=["GET", "POST"])
def update_settings():
    global settings
    if request.method == "POST":
        data = request.json
        settings.update({k: v for k, v in data.items() if k in settings})
        translation_queue.put(("__settings__", settings))
    return jsonify(settings)


# Cloudflare 등 프록시의 버퍼링을 깨기 위한 초기 패딩(2KB 주석)
_SSE_PADDING = ":" + (" " * 2048) + "\n\n"


@sock.route("/ws-display")
def ws_display(ws):
    """자막 송출용 WebSocket (터널 환경에서 SSE 버퍼링 회피). 송출창+셀폰 공용."""
    q = translation_queue.subscribe()
    try:
        ws.send(json.dumps({"type": "settings", "data": settings}))
        while True:
            try:
                item = q.get(timeout=10)
            except queue.Empty:
                ws.send(json.dumps({"type": "ping"}))
                continue
            if isinstance(item, tuple) and item[0] == "__settings__":
                msg = {"type": "settings", "data": item[1]}
            elif item == "__clear__":
                msg = {"type": "clear"}
            elif item == "__done__":
                msg = {"type": "done"}
            else:
                msg = {"type": "chunk", "data": item}
            ws.send(json.dumps(msg))
    except Exception:
        pass
    finally:
        translation_queue.unsubscribe(q)


@app.route("/stream")
def stream():
    q = translation_queue.subscribe()

    def event_stream():
        yield _SSE_PADDING  # 버퍼 강제 비우기
        try:
            while True:
                try:
                    item = q.get(timeout=1)
                except queue.Empty:
                    yield ": ping\n\n"  # 하트비트(연결 유지 + 버퍼 flush)
                    continue
                if isinstance(item, tuple) and item[0] == "__settings__":
                    yield f"data: __settings__{json.dumps(item[1])}\n\n"
                elif item == "__clear__":
                    yield "data: __clear__\n\n"
                elif item == "__done__":
                    yield "data: __done__\n\n"
                else:
                    escaped = item.replace("\n", "\\n")
                    yield f"data: {escaped}\n\n"
        finally:
            translation_queue.unsubscribe(q)

    return Response(event_stream(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route("/operator-stream")
def operator_stream():
    q = operator_queue.subscribe()

    def event_stream():
        yield _SSE_PADDING
        yield f"data: {json.dumps({'type': 'settings', 'data': settings})}\n\n"
        try:
            while True:
                try:
                    event_type, data = q.get(timeout=1)
                except queue.Empty:
                    yield ": ping\n\n"
                    continue
                yield f"data: {json.dumps({'type': event_type, 'data': data})}\n\n"
        finally:
            operator_queue.unsubscribe(q)

    return Response(event_stream(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


if __name__ == "__main__":
    print("서버 시작")
    print("  운영자 대시보드: http://localhost:5000/operator")
    print("  출력 화면:       http://localhost:5000")
    app.run(host="0.0.0.0", port=5000, threaded=True)
