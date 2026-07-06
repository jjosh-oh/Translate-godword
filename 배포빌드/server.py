"""
Translation server with operator dashboard.
"""
import os
import sys
import io
import json
import queue
import threading
import anthropic
from flask import Flask, request, jsonify, Response, send_from_directory
from flask_sock import Sock
from dotenv import load_dotenv

# 실행 위치 판별: .exe로 묶였을 때와 일반 실행을 모두 지원
#  - BUNDLE_DIR: 화면 파일 등 '내장' 자원 위치 (읽기 전용)
#  - APP_DIR   : 사용자 파일 위치 (.env, google-key.json, 로그 등)
#                .exe 실행 시 → .exe 폴더 (쓰기 불가하면 %APPDATA%\LiveWord)
#                개발 실행 시 → 스크립트 폴더
if getattr(sys, "frozen", False):
    BUNDLE_DIR = sys._MEIPASS
    # 1순위: .exe 폴더 (설치 위치가 사용자 폴더라 대부분 쓰기 가능,
    #         AppData는 Windows 앱 격리로 접근이 차단되는 경우가 있음)
    # 2순위: %APPDATA%\LiveWord
    _exe_dir = os.path.dirname(sys.executable)
    try:
        _t = os.path.join(_exe_dir, ".write_test")
        with open(_t, "w") as _f:
            _f.write("ok")
        os.remove(_t)
        APP_DIR = _exe_dir
    except Exception:
        APP_DIR = os.path.join(os.environ.get("APPDATA", _exe_dir),
                               "LiveWord")
        os.makedirs(APP_DIR, exist_ok=True)
else:
    BUNDLE_DIR = os.path.dirname(os.path.abspath(__file__))
    APP_DIR = BUNDLE_DIR

# 설정(.env)은 APP_DIR에서 읽음
load_dotenv(os.path.join(APP_DIR, ".env"), override=True, encoding="utf-8-sig")

# Google Cloud 인증: .exe 옆 google-key.json 자동 사용
_key_path = os.path.join(APP_DIR, "google-key.json")
if os.path.exists(_key_path) and "GOOGLE_APPLICATION_CREDENTIALS" not in os.environ:
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = _key_path

app = Flask(__name__)
sock = Sock(app)

# ── ngrok 자동 시작 ─────────────────────────────────────────────────────────
_ngrok_url = ""  # 실제 연결된 공개 URL (터널이 뜨면 채워짐)

def _start_ngrok():
    """백그라운드에서 ngrok 터널을 시작하고 URL을 _ngrok_url에 저장."""
    global _ngrok_url
    import subprocess, time, re

    authtoken = os.environ.get("NGROK_AUTHTOKEN", "").strip()
    domain = os.environ.get("NGROK_DOMAIN", "").strip()
    if not authtoken:
        return  # 설정 안 됐으면 건너뜀

    # ngrok 실행 파일: .exe 폴더에 번들된 것 우선, 없으면 시스템 PATH
    ngrok_exe = "ngrok"
    if getattr(sys, "frozen", False):
        _bundled = os.path.join(os.path.dirname(sys.executable), "ngrok.exe")
        if os.path.exists(_bundled):
            ngrok_exe = _bundled

    # 고정 도메인이 있으면 URL을 미리 설정 (파싱 기다릴 필요 없음)
    if domain:
        _ngrok_url = "https://" + domain
        print(f"[ngrok] 고정 도메인: {_ngrok_url}")

    # authtoken 등록 (최초 1회 또는 변경 시)
    try:
        subprocess.run([ngrok_exe, "config", "add-authtoken", authtoken],
                       capture_output=True, timeout=10)
    except Exception:
        return

    # 터널 시작
    cmd = [ngrok_exe, "http", "5000", "--log=stdout", "--log-format=json"]
    if domain:
        cmd += ["--domain", domain]
    try:
        proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                                text=True, encoding="utf-8", errors="ignore")
    except FileNotFoundError:
        print("[ngrok] ngrok을 찾을 수 없습니다. 설치 여부를 확인하세요.")
        return

    # stdout에서 URL 파싱 (고정 도메인 없는 경우 여기서 URL 확보)
    if not domain:
        for line in proc.stdout:
            try:
                obj = json.loads(line)
                url = obj.get("url") or obj.get("public_url") or ""
                if url.startswith("https://"):
                    _ngrok_url = url
                    print(f"[ngrok] 터널 연결: {_ngrok_url}")
                    break
            except Exception:
                if "url=" in line:
                    m = re.search(r"url=(https://\S+)", line)
                    if m:
                        _ngrok_url = m.group(1)
                        print(f"[ngrok] 터널 연결: {_ngrok_url}")
                        break

if os.environ.get("NGROK_AUTHTOKEN"):
    threading.Thread(target=_start_ngrok, daemon=True).start()

# Anthropic 클라이언트: 키가 없어도 서버는 떠야 하므로 지연 생성
client = None
def get_client():
    global client
    if client is None:
        client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
    return client

# Gemini 클라이언트(비교용) — Vertex AI 경유, 실패해도 서버는 정상 동작
GEMINI_PROJECT = "project-71ecd8e1-9699-417f-ae4"
gemini_client = None
try:
    from google import genai as _genai
    gemini_client = _genai.Client(vertexai=True, project=GEMINI_PROJECT, location="us-central1")
except Exception as _e:
    print("Gemini 초기화 건너뜀:", _e)

# 음성인식 언어 코드 매핑 (Google Speech 형식)
STT_LANG = {
    "ko-KR": "ko-KR", "en-US": "en-US", "ja-JP": "ja-JP",
    "zh-CN": "cmn-Hans-CN", "es-ES": "es-ES", "fr-FR": "fr-FR", "de-DE": "de-DE",
}
STT_TO_NAME = {
    "ko-KR": "Korean", "en-US": "English", "ja-JP": "Japanese",
    "zh-CN": "Chinese", "es-ES": "Spanish", "fr-FR": "French", "de-DE": "German",
}

# 여러 기기(송출창 + 셀폰들)에 동시 전송하기 위한 구독자 기반 브로드캐스트
_display_subscribers = set()
_operator_subscribers = set()
_sub_lock = threading.Lock()


class _Broadcaster:
    """put() 호출 시 모든 구독자 큐에 메시지를 복사해 넣는다."""
    def __init__(self, subscribers):
        self._subs = subscribers

    def put(self, item):
        with _sub_lock:
            for q in list(self._subs):
                q.put(item)

    def subscribe(self):
        q = queue.Queue()
        with _sub_lock:
            self._subs.add(q)
        return q

    def unsubscribe(self, q):
        with _sub_lock:
            self._subs.discard(q)


translation_queue = _Broadcaster(_display_subscribers)
operator_queue = _Broadcaster(_operator_subscribers)

# Shared display settings
settings = {
    "font_size": 5,
    "text_color": "#ffffff",
    "bg_color": "#000000",
    "delay": 0,
    "target_lang": "English",
    "source_lang": "Korean",
}

last_input = ""
last_output = ""
sermon_context = ""  # 업로드된 설교 자료

# 교회 용어집(인식 힌트 전용) — glossary.txt에서 한 줄에 하나씩 로드
glossary_terms = []


def load_glossary():
    global glossary_terms
    # .exe 옆 사용자 용어집을 우선, 없으면 내장 기본 용어집
    path = os.path.join(APP_DIR, "glossary.txt")
    if not os.path.exists(path):
        path = os.path.join(BUNDLE_DIR, "glossary.txt")
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            lines = [ln.strip() for ln in f.read().splitlines()]
        glossary_terms = [ln for ln in lines if ln]
    except Exception:
        glossary_terms = []


load_glossary()

# 번역 대응표 (한글=English 형식)
translation_mapping = {}

def load_mapping():
    global translation_mapping
    path = os.path.join(APP_DIR, "mapping.txt")
    try:
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    line = line.strip()
                    if "=" in line and line:
                        k, v = line.split("=", 1)
                        translation_mapping[k.strip()] = v.strip()
    except Exception:
        translation_mapping = {}

load_mapping()


def _log_translation(src, out):
    try:
        path = os.path.join(APP_DIR, "로그.txt")
        import datetime
        ts = datetime.datetime.now().strftime("%H:%M:%S")
        with open(path, "a", encoding="utf-8") as f:
            f.write(f"[{ts}] 입력: {src}\n[{ts}] 번역: {out}\n\n")
    except Exception:
        pass


def translate_and_stream(text: str, target_lang: str, source_lang: str):
    global last_output
    base_instruction = (
        f"You are a professional church interpreter providing live subtitles. "
        f"Translate ONLY the exact given text from {source_lang} into {target_lang}. "
        f"Output ONLY the translation of what is given — never add, continue, complete, "
        f"or quote additional text. If the input is a short or partial sentence (e.g. part of "
        f"a Bible verse), translate only that fragment; do NOT finish the verse or add the rest. "
        f"No explanations or commentary."
    )

    # 번역 대응표가 있으면 프롬프트에 추가
    mapping_text = ""
    if translation_mapping:
        pairs = "\n".join(f"  {k} → {v}" for k, v in translation_mapping.items())
        mapping_text = (
            "\n\nName/term translation table (use these EXACT translations when the term appears):\n"
            + pairs
        )

    if sermon_context:
        # 설교 원고: 고유명사 확인 '참고용'으로만 — 내용을 이어 쓰지 않도록 명시
        reference = (
            "The sermon script below is a REFERENCE ONLY, used solely to spell proper nouns "
            "(biblical names, place names, Korean locations) consistently. "
            "Do NOT translate, output, continue, or copy any sentence from this script. "
            "Only translate the user's given text.\n\n"
            "--- SERMON REFERENCE (do not output) ---\n" + sermon_context[:8000] + "\n--- END ---"
        )
        system = [
            {"type": "text", "text": base_instruction + mapping_text},
            {"type": "text", "text": reference,
             "cache_control": {"type": "ephemeral"}},
        ]
    else:
        system = base_instruction + mapping_text

    translation_queue.put("__clear__")
    operator_queue.put(("output_clear", ""))

    # 출력 폭주 방지: 입력 길이에 비례한 상한(최소 256, 최대 1024 토큰)
    max_out = max(256, min(1024, len(text) * 4))

    result = []
    with get_client().messages.stream(
        model="claude-opus-4-8",
        max_tokens=max_out,
        system=system,
        messages=[{"role": "user", "content": text}],
    ) as stream:
        for chunk in stream.text_stream:
            result.append(chunk)
            translation_queue.put(chunk)
            operator_queue.put(("output_chunk", chunk))

    last_output = "".join(result)
    _log_translation(text, last_output)
    translation_queue.put("__done__")
    operator_queue.put(("done", ""))


# ===== 번역 작업 큐 (문장 단위 번역을 순서대로 처리) =====
translation_jobs = queue.Queue()

# 음성 세션 관리: 마이크를 끄면 세션이 바뀌어 대기 중 번역이 폐기됨
_current_session = 0
_session_lock = threading.Lock()


def _translation_worker():
    while True:
        text, target_lang, source_lang, session = translation_jobs.get()
        # 세션이 유효한 경우에만 번역 (None이면 수동 입력 → 항상 실행)
        if session is not None:
            with _session_lock:
                if session != _current_session:
                    continue  # 마이크가 꺼진 뒤의 오래된 작업 → 폐기
        try:
            translate_and_stream(text, target_lang, source_lang)
        except Exception as e:
            print("번역 오류:", e)


threading.Thread(target=_translation_worker, daemon=True).start()


def enqueue_translation(text, target_lang, source_lang, session=None):
    translation_jobs.put((text, target_lang, source_lang, session))


@app.route("/")
def index():
    return send_from_directory(BUNDLE_DIR, "display.html")


@app.route("/operator")
def operator():
    return send_from_directory(BUNDLE_DIR, "operator.html")


@app.route("/mobile")
def mobile():
    return send_from_directory(BUNDLE_DIR, "mobile.html")


@app.route("/compare")
def compare():
    return send_from_directory(BUNDLE_DIR, "compare.html")


def _translate_claude(text, target_lang, source_lang):
    system = (f"You are a professional church interpreter. Translate from {source_lang} into {target_lang}. "
              f"Output ONLY the translation.")
    m = get_client().messages.create(model="claude-opus-4-8", max_tokens=1024,
                               system=system, messages=[{"role": "user", "content": text}])
    return "".join(b.text for b in m.content if getattr(b, "type", "") == "text")


def _translate_gemini(text, target_lang, source_lang):
    if not gemini_client:
        return "(Gemini 사용 불가)"
    prompt = (f"You are a professional church interpreter. Translate from {source_lang} into {target_lang}. "
              f"Output ONLY the translation.\n\n{text}")
    r = gemini_client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
    return (r.text or "").strip()


@sock.route("/audio-compare")
def audio_compare(ws):
    """비교용: 음성인식 후 같은 문장을 Claude·Gemini로 번역해 나란히 전송."""
    from google.cloud import speech
    import time as _t

    cfg = json.loads(ws.receive())
    src_code = STT_LANG.get(cfg.get("source_lang_code"), "ko-KR")
    source_name = STT_TO_NAME.get(cfg.get("source_lang_code"), "Korean")
    target_lang = cfg.get("target_lang", "English")

    speech_client = speech.SpeechClient()
    recog_config = speech.RecognitionConfig(
        encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=16000, language_code=src_code,
        enable_automatic_punctuation=True, model="latest_long", use_enhanced=True,
    )
    streaming_config = speech.StreamingRecognitionConfig(config=recog_config, interim_results=True)

    audio_q = queue.Queue()
    stop_flag = {"stop": False}
    send_lock = threading.Lock()

    def safe_send(obj):
        try:
            with send_lock:
                ws.send(json.dumps(obj))
        except Exception:
            pass

    def receiver():
        try:
            while True:
                data = ws.receive()
                if data is None:
                    break
                if isinstance(data, (bytes, bytearray)):
                    audio_q.put(bytes(data))
        except Exception:
            pass
        finally:
            stop_flag["stop"] = True
            audio_q.put(None)

    threading.Thread(target=receiver, daemon=True).start()

    def request_gen():
        while not stop_flag["stop"]:
            chunk = audio_q.get()
            if chunk is None:
                return
            yield speech.StreamingRecognizeRequest(audio_content=chunk)

    def run_engine(engine, fn, text):
        t0 = _t.time()
        try:
            out = fn(text, target_lang, source_name)
        except Exception as e:
            out = "(오류: " + str(e)[:60] + ")"
        ms = int((_t.time() - t0) * 1000)
        safe_send({"engine": engine, "text": out, "ms": ms})

    import re as _rec

    def split_sentences(t):
        return [p for p in _rec.split(r'(?<=[\.\?\!。？！…])\s*', t) if p.strip()]

    def translate_both(text):
        safe_send({"engine": "input", "text": text})
        threading.Thread(target=run_engine, args=("claude", _translate_claude, text), daemon=True).start()
        threading.Thread(target=run_engine, args=("gemini", _translate_gemini, text), daemon=True).start()

    translated = []

    while not stop_flag["stop"]:
        try:
            responses = speech_client.streaming_recognize(streaming_config, request_gen())
            for response in responses:
                for result in response.results:
                    tr = result.alternatives[0].transcript
                    if not tr.strip():
                        continue
                    pieces = split_sentences(tr)
                    if result.is_final:
                        for p in pieces:
                            n = p.strip()
                            if n and n not in translated:
                                translate_both(n)
                        translated = []
                    else:
                        safe_send({"engine": "interim", "text": tr.strip()})
                        if len(pieces) >= 2:
                            for p in pieces[:-1]:
                                n = p.strip()
                                if n and n not in translated:
                                    translated.append(n)
                                    translate_both(n)
        except Exception as e:
            safe_send({"engine": "error", "text": str(e)[:120]})
            if stop_flag["stop"]:
                break


@app.route("/tunnel-url")
def tunnel_url():
    """공개 주소 반환. 우선순위: 자동 ngrok > ngrok-domain.txt > cloudflare."""
    import re
    base = APP_DIR

    # 1) 자동 시작된 ngrok URL
    if _ngrok_url:
        return jsonify({"url": _ngrok_url})

    # 2) 수동 설정된 ngrok 고정 도메인 파일
    domain_path = os.path.join(base, "ngrok-domain.txt")
    try:
        if os.path.exists(domain_path):
            with open(domain_path, "r", encoding="utf-8", errors="ignore") as f:
                domain = f.read().strip()
            if domain:
                return jsonify({"url": "https://" + domain})
    except Exception:
        pass

    # 3) cloudflare 임시 주소(tunnel.log)
    log_path = os.path.join(base, "tunnel.log")
    try:
        with open(log_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        matches = re.findall(r"https://[a-zA-Z0-9-]+\.trycloudflare\.com", text)
        if matches:
            return jsonify({"url": matches[-1]})
    except Exception:
        pass
    return jsonify({"url": ""})


@sock.route("/audio")
def audio_socket(ws):
    """브라우저에서 16kHz PCM 오디오를 받아 Google Speech로 스트리밍 인식 → 번역."""
    from google.cloud import speech

    # 첫 메시지: 설정(JSON)
    cfg_raw = ws.receive()
    cfg = json.loads(cfg_raw)
    src_code = STT_LANG.get(cfg.get("source_lang_code"), "ko-KR")
    source_name = STT_TO_NAME.get(cfg.get("source_lang_code"), "Korean")
    target_lang = cfg.get("target_lang", "English")

    # 인식 힌트(speech adaptation)
    speech_contexts = []
    # 1) 교회 용어집 — 모든 항목을 높은 가중치로(고유명사 사전)
    if glossary_terms:
        speech_contexts.append(speech.SpeechContext(phrases=glossary_terms[:4000], boost=20.0))
    # 2) 오늘 설교 원고에서 자주 나오는 단어 — 보조 힌트
    if sermon_context:
        import re as _re2
        from collections import Counter
        words = _re2.findall(r"[가-힣]{2,}", sermon_context)
        common = [w for w, _ in Counter(words).most_common(300)]
        if common:
            speech_contexts.append(speech.SpeechContext(phrases=common, boost=12.0))

    speech_client = speech.SpeechClient()
    recog_config = speech.RecognitionConfig(
        encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=16000,
        language_code=src_code,
        enable_automatic_punctuation=True,
        model="latest_long",        # 긴 발화에 적합한 최신 모델
        use_enhanced=True,          # 고품질(enhanced) 모델 사용
        speech_contexts=speech_contexts,
    )
    streaming_config = speech.StreamingRecognitionConfig(
        config=recog_config, interim_results=True
    )

    audio_q = queue.Queue()
    stop_flag = {"stop": False}

    # 브라우저 → 오디오 수신 스레드
    def receiver():
        try:
            while True:
                data = ws.receive()
                if data is None:
                    break
                if isinstance(data, (bytes, bytearray)):
                    audio_q.put(bytes(data))
        except Exception:
            pass
        finally:
            stop_flag["stop"] = True
            audio_q.put(None)

    recv_thread = threading.Thread(target=receiver, daemon=True)
    recv_thread.start()

    def request_gen():
        # Google 스트림은 약 5분 제한 → 호출부에서 재시작
        while not stop_flag["stop"]:
            chunk = audio_q.get()
            if chunk is None:
                return
            yield speech.StreamingRecognizeRequest(audio_content=chunk)

    # 이 연결의 세션 ID (마이크를 끄면 무효화되어 대기 중 번역이 폐기됨)
    global _current_session
    with _session_lock:
        _current_session += 1
        my_session = _current_session

    import re as _re

    def split_sentences(t):
        # 문장부호(. ? ! 。 ？ ！ …) 뒤에서 자른다. 부호를 포함해 반환.
        parts = _re.split(r'(?<=[\.\?\!。？！…])\s*', t)
        return [p for p in parts if p.strip()]

    translated = []  # 현재 발화에서 이미 번역에 보낸 문장(텍스트로 중복 제거)

    # 5분 제한 대응: 스트림이 끝나면 다시 연결 (음성 큐는 유지)
    while not stop_flag["stop"]:
        try:
            responses = speech_client.streaming_recognize(streaming_config, request_gen())
            for response in responses:
                for result in response.results:
                    transcript = result.alternatives[0].transcript
                    if not transcript.strip():
                        continue
                    pieces = split_sentences(transcript)

                    if result.is_final:
                        operator_queue.put(("input", transcript.strip()))
                        # 아직 안 보낸 문장(끝 미완성 조각 포함)을 모두 번역
                        for p in pieces:
                            n = p.strip()
                            if n and n not in translated:
                                enqueue_translation(n, target_lang, source_name, my_session)
                        translated = []  # 다음 발화 위해 초기화
                    else:
                        operator_queue.put(("interim", transcript.strip()))
                        # 마침표로 '완성된' 문장만 즉시 번역 (마지막 미완성 조각 제외)
                        if len(pieces) >= 2:
                            complete = pieces[:-1]
                            for p in complete:
                                n = p.strip()
                                if n and n not in translated:
                                    translated.append(n)
                                    enqueue_translation(n, target_lang, source_name, my_session)
        except Exception as e:
            operator_queue.put(("stt_error", str(e)[:120]))
            if stop_flag["stop"]:
                break

    # 마이크 종료: 이 세션을 무효화해 대기 중인 번역 작업을 폐기
    with _session_lock:
        if _current_session == my_session:
            _current_session += 1


@app.route("/translate", methods=["POST"])
def translate():
    global last_input
    data = request.json
    text = data.get("text", "").strip()
    target_lang = data.get("target_lang", settings["target_lang"])
    source_lang = data.get("source_lang", settings["source_lang"])

    if not text:
        return jsonify({"error": "텍스트를 입력하세요."}), 400

    last_input = text
    operator_queue.put(("input", text))

    thread = threading.Thread(target=translate_and_stream, args=(text, target_lang, source_lang))
    thread.daemon = True
    thread.start()

    return jsonify({"status": "started"})


@app.route("/upload", methods=["POST"])
def upload():
    global sermon_context
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "파일이 없습니다."}), 400

    filename = file.filename.lower()
    try:
        if filename.endswith(".txt"):
            sermon_context = file.read().decode("utf-8", errors="ignore")

        elif filename.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file.read()))
            sermon_context = "\n".join(page.extract_text() or "" for page in reader.pages)

        elif filename.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(file.read()))
            sermon_context = "\n".join(p.text for p in doc.paragraphs)

        else:
            return jsonify({"error": "지원 형식: .txt, .pdf, .docx"}), 400

        preview = sermon_context[:200].replace("\n", " ")
        return jsonify({"status": "ok", "chars": len(sermon_context), "preview": preview})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/upload-clear", methods=["POST"])
def upload_clear():
    global sermon_context
    sermon_context = ""
    return jsonify({"status": "cleared"})


def _extract_text(file):
    name = file.filename.lower()
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file.read()))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if name.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(file.read()))
        return "\n".join(p.text for p in doc.paragraphs)
    return file.read().decode("utf-8", errors="ignore")


@app.route("/upload-glossary", methods=["POST"])
def upload_glossary():
    """교회 용어집 업로드 — 기존 용어집에 병합(중복 제거)하여 저장."""
    global glossary_terms
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "파일이 없습니다."}), 400
    try:
        text = _extract_text(file)
        # 기존 용어 유지 + 새 용어 추가 (순서 보존, 중복 제거)
        seen = set(glossary_terms)
        merged = list(glossary_terms)
        new_count = 0
        for ln in text.splitlines():
            t = ln.strip()
            if t and t not in seen:
                seen.add(t)
                merged.append(t)
                new_count += 1
        glossary_terms = merged
        # glossary.txt에 저장(다음 실행에도 유지)
        path = os.path.join(APP_DIR, "glossary.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(merged))
        return jsonify({"status": "ok", "count": len(merged), "added": new_count,
                        "preview": ", ".join(merged[:15])})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/glossary-info")
def glossary_info():
    return jsonify({"count": len(glossary_terms),
                    "preview": ", ".join(glossary_terms[:15])})


@app.route("/upload-mapping", methods=["POST"])
def upload_mapping():
    """번역 대응표 업로드 — 한글=English 형식, 기존 항목에 병합."""
    global translation_mapping
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "파일이 없습니다."}), 400
    try:
        text = _extract_text(file)
        new_count = 0
        for line in text.splitlines():
            line = line.strip()
            if "=" in line and line:
                k, v = line.split("=", 1)
                k, v = k.strip(), v.strip()
                if k and v:
                    if k not in translation_mapping:
                        new_count += 1
                    translation_mapping[k] = v
        path = os.path.join(APP_DIR, "mapping.txt")
        with open(path, "w", encoding="utf-8") as f:
            for k, v in translation_mapping.items():
                f.write(f"{k}={v}\n")
        preview = ", ".join(f"{k}→{v}" for k, v in list(translation_mapping.items())[:5])
        return jsonify({"status": "ok", "count": len(translation_mapping),
                        "added": new_count, "preview": preview})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/mapping-info")
def mapping_info():
    preview = ", ".join(f"{k}→{v}" for k, v in list(translation_mapping.items())[:5])
    return jsonify({"count": len(translation_mapping), "preview": preview})


@app.route("/settings", methods=["GET", "POST"])
def update_settings():
    global settings
    if request.method == "POST":
        data = request.json
        settings.update({k: v for k, v in data.items() if k in settings})
        translation_queue.put(("__settings__", settings))
    return jsonify(settings)


# Cloudflare 등 프록시의 버퍼링을 깨기 위한 초기 패딩(2KB 주석)
_SSE_PADDING = ":" + (" " * 2048) + "\n\n"


@sock.route("/ws-display")
def ws_display(ws):
    """자막 송출용 WebSocket (터널 환경에서 SSE 버퍼링 회피). 송출창+셀폰 공용."""
    q = translation_queue.subscribe()
    try:
        ws.send(json.dumps({"type": "settings", "data": settings}))
        while True:
            try:
                item = q.get(timeout=10)
            except queue.Empty:
                ws.send(json.dumps({"type": "ping"}))
                continue
            if isinstance(item, tuple) and item[0] == "__settings__":
                msg = {"type": "settings", "data": item[1]}
            elif item == "__clear__":
                msg = {"type": "clear"}
            elif item == "__done__":
                msg = {"type": "done"}
            else:
                msg = {"type": "chunk", "data": item}
            ws.send(json.dumps(msg))
    except Exception:
        pass
    finally:
        translation_queue.unsubscribe(q)


@app.route("/stream")
def stream():
    q = translation_queue.subscribe()

    def event_stream():
        yield _SSE_PADDING  # 버퍼 강제 비우기
        try:
            while True:
                try:
                    item = q.get(timeout=1)
                except queue.Empty:
                    yield ": ping\n\n"  # 하트비트(연결 유지 + 버퍼 flush)
                    continue
                if isinstance(item, tuple) and item[0] == "__settings__":
                    yield f"data: __settings__{json.dumps(item[1])}\n\n"
                elif item == "__clear__":
                    yield "data: __clear__\n\n"
                elif item == "__done__":
                    yield "data: __done__\n\n"
                else:
                    escaped = item.replace("\n", "\\n")
                    yield f"data: {escaped}\n\n"
        finally:
            translation_queue.unsubscribe(q)

    return Response(event_stream(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route("/operator-stream")
def operator_stream():
    q = operator_queue.subscribe()

    def event_stream():
        yield _SSE_PADDING
        yield f"data: {json.dumps({'type': 'settings', 'data': settings})}\n\n"
        try:
            while True:
                try:
                    event_type, data = q.get(timeout=1)
                except queue.Empty:
                    yield ": ping\n\n"
                    continue
                yield f"data: {json.dumps({'type': event_type, 'data': data})}\n\n"
        finally:
            operator_queue.unsubscribe(q)

    return Response(event_stream(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ── 설정 화면 (첫 실행 시 API 키 입력) ───────────────────────────────────────

@app.route("/setup")
def setup():
    return send_from_directory(BUNDLE_DIR, "setup.html")


@app.route("/setup-status")
def setup_status():
    has_anthropic = bool(os.environ.get("ANTHROPIC_API_KEY"))
    has_ngrok_token = bool(os.environ.get("NGROK_AUTHTOKEN"))
    ngrok_domain = os.environ.get("NGROK_DOMAIN", "")
    if not has_anthropic or not has_ngrok_token:
        env_path = os.path.join(APP_DIR, ".env")
        if os.path.exists(env_path):
            with open(env_path, "r", encoding="utf-8-sig") as f:
                content = f.read()
            if not has_anthropic:
                has_anthropic = "ANTHROPIC_API_KEY=" in content
            if not has_ngrok_token:
                has_ngrok_token = "NGROK_AUTHTOKEN=" in content
    has_google = os.path.exists(os.path.join(APP_DIR, "google-key.json"))
    return jsonify(
        has_anthropic_key=has_anthropic,
        has_google_key=has_google,
        has_ngrok_token=has_ngrok_token,
        ngrok_domain=ngrok_domain,
        ngrok_url=_ngrok_url,
    )


@app.route("/setup-save", methods=["POST"])
def setup_save():
    data = request.get_json(force=True)
    anthropic_key = (data.get("anthropic_key") or "").strip()
    ngrok_token   = (data.get("ngrok_token") or "").strip()
    ngrok_domain  = (data.get("ngrok_domain") or "").strip()

    if not anthropic_key and not ngrok_token and not ngrok_domain:
        return jsonify(ok=False, error="저장할 값이 없습니다")
    if anthropic_key and not anthropic_key.startswith("sk-"):
        return jsonify(ok=False, error="올바른 Anthropic 키 형식이 아닙니다 (sk-ant-... 로 시작해야 함)")

    try:
        env_path = os.path.join(APP_DIR, ".env")
        # 기존 .env 읽어서 해당 키만 교체
        env_vars = {}
        if os.path.exists(env_path):
            with open(env_path, "r", encoding="utf-8-sig") as f:
                for line in f:
                    line = line.rstrip("\n")
                    if "=" in line and not line.startswith("#"):
                        k, v = line.split("=", 1)
                        env_vars[k.strip()] = v.strip()
        if anthropic_key:
            env_vars["ANTHROPIC_API_KEY"] = anthropic_key
            os.environ["ANTHROPIC_API_KEY"] = anthropic_key
            global client
            client = None
        if ngrok_token:
            env_vars["NGROK_AUTHTOKEN"] = ngrok_token
            os.environ["NGROK_AUTHTOKEN"] = ngrok_token
        if ngrok_domain:
            env_vars["NGROK_DOMAIN"] = ngrok_domain
            os.environ["NGROK_DOMAIN"] = ngrok_domain

        with open(env_path, "w", encoding="utf-8") as f:
            for k, v in env_vars.items():
                f.write(f"{k}={v}\n")

        # ngrok 토큰이 새로 저장됐으면 터널 시작
        if ngrok_token and not _ngrok_url:
            threading.Thread(target=_start_ngrok, daemon=True).start()

        return jsonify(ok=True)
    except Exception as e:
        return jsonify(ok=False, error=str(e))


@app.route("/setup-upload-google-key", methods=["POST"])
def setup_upload_google_key():
    f = request.files.get("file")
    if not f:
        return jsonify(ok=False, error="파일이 없습니다")
    try:
        content = f.read()
        parsed = json.loads(content)
        if parsed.get("type") != "service_account":
            return jsonify(ok=False, error="서비스 계정 JSON 파일이 아닙니다")
        dest = os.path.join(APP_DIR, "google-key.json")
        with open(dest, "wb") as out:
            out.write(content)
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = dest
        return jsonify(ok=True)
    except json.JSONDecodeError:
        return jsonify(ok=False, error="JSON 형식이 올바르지 않습니다")
    except Exception as e:
        return jsonify(ok=False, error=str(e))


@app.route("/setup-test-anthropic", methods=["POST"])
def setup_test_anthropic():
    data = request.get_json(force=True)
    key = (data.get("api_key") or "").strip() or os.environ.get("ANTHROPIC_API_KEY", "")
    if not key:
        return jsonify(ok=False, error="API 키를 입력해 주세요")
    try:
        import anthropic as _anthropic
        test_client = _anthropic.Anthropic(api_key=key)
        msg = test_client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=10,
            messages=[{"role": "user", "content": "Hi"}]
        )
        return jsonify(ok=True, model=msg.model)
    except Exception as e:
        return jsonify(ok=False, error=str(e))


@app.route("/setup-test-google", methods=["POST"])
def setup_test_google():
    try:
        from google.cloud import speech
        client_stt = speech.SpeechClient()
        # 간단한 인식 요청으로 인증 확인 (빈 오디오는 오류지만 인증 성공 여부만 체크)
        client_stt.recognize(
            config=speech.RecognitionConfig(
                encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
                sample_rate_hertz=16000,
                language_code="ko-KR",
            ),
            audio=speech.RecognitionAudio(content=b"\x00" * 32000),
        )
        return jsonify(ok=True)
    except Exception as e:
        err = str(e)
        # 빈 오디오로 인한 오류는 인증 성공으로 간주
        if "no speech" in err.lower() or "audio" in err.lower() or "empty" in err.lower():
            return jsonify(ok=True)
        if "credentials" in err.lower() or "permission" in err.lower() or "auth" in err.lower():
            return jsonify(ok=False, error="인증 실패: " + err)
        # 그 외 오류는 연결 자체는 된 것
        return jsonify(ok=True)


def _open_browser():
    """서버가 뜬 뒤 브라우저를 자동으로 연다. 설정이 안 됐으면 /setup, 됐으면 /operator."""
    import time, webbrowser
    time.sleep(1.5)  # 서버 기동 대기
    has_key = bool(os.environ.get("ANTHROPIC_API_KEY"))
    url = "http://localhost:5000/" + ("operator" if has_key else "setup")
    print(f"  브라우저 열기: {url}")
    webbrowser.open(url)


def _already_running():
    """포트 5000에 이미 서버가 떠 있는지 확인."""
    import socket
    try:
        s = socket.create_connection(("127.0.0.1", 5000), timeout=1)
        s.close()
        return True
    except Exception:
        return False


if __name__ == "__main__":
    if _already_running():
        import webbrowser
        print("프로그램이 이미 실행 중입니다. 기존 화면을 엽니다.")
        webbrowser.open("http://localhost:5000/operator")
        sys.exit(0)
    print("=" * 50)
    print("  LiveWord 서버 시작")
    print(f"  설정 파일 위치: {APP_DIR}")
    print("  설정 화면:       http://localhost:5000/setup")
    print("  운영자 대시보드: http://localhost:5000/operator")
    print("=" * 50)
    threading.Thread(target=_open_browser, daemon=True).start()
    app.run(host="0.0.0.0", port=5000, threaded=True)
